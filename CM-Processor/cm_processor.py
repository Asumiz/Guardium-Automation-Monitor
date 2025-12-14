"""
Guardium CM Processor — Health Check Completo
===================================================

Este script automatiza o Health Check de ambientes Guardium, processando logs e gerando relatórios
detalhados em Word e Excel.

Estrutura de Repositório Esperada:
/
├── CM-Processor/
│   └── cm_processor.py <--- (Este arquivo)
└── CM/ <--- (Pasta de trabalho criada na raiz)
    ├── Central Management/
    ├── STAP status/
    ...

✔ Estrutura de pastas limpa (Sem "Processos Internos" ou "Tabelas Internas")
✔ Limpeza automática de arquivos de execuções anteriores
✔ STAP:
    - Contagem Ativos vs Inativos
    - Lista detalhada no Word dos INATIVOS (Host + Versão)
✔ Agregação:
    - Detecta falhas (Purge, Archive, Export)
    - Captura a DATA da falha
    - Relata no Word: Coletor - Falha - Data

Dependências:
    pip install pandas openpyxl python-docx
"""

import sys
from pathlib import Path
from datetime import datetime
import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except Exception:
    Document = None

# =========================================================
# CONFIGURAÇÃO
# =========================================================

# Apenas estas pastas serão criadas
BASE_SUBFOLDERS = [
    "Central Management",
    "STAP status",
    "Processos de agregação",
    "Qualidade da coleta",
]

ACTIVE_KEYWORDS = ("active", "up", "running", "connected", "online")
INACTIVE_KEYWORDS = ("inactive", "down", "stopped", "disconnected", "offline", "failed", "error")
SUCCESS_KEYWORDS = ("success", "done", "completed", "ok")

# =========================================================
# FUNÇÕES AUXILIARES
# =========================================================

def read_table(path: Path) -> pd.DataFrame:
    try:
        if path.suffix.lower() in (".xls", ".xlsx"):
            return pd.read_excel(path)
        if path.suffix.lower() == ".csv":
            return pd.read_csv(path)
    except Exception as e:
        print(f"⚠ Erro ao ler {path.name}: {e}")
    return pd.DataFrame()

def find_column(df, keywords):
    """Encontra uma coluna baseada em palavras-chave (case insensitive)"""
    for c in df.columns:
        c_str = str(c).lower()
        for k in keywords:
            if k in c_str:
                return c
    return None

def clean_folder_contents(folder: Path, recursive=False):
    """Limpa arquivos anteriores para evitar dados duplicados/antigos"""
    if not folder.is_dir():
        return
    
    for item in folder.iterdir():
        if item.is_file():
            try:
                item.unlink()
            except Exception:
                pass
        elif recursive and item.is_dir():
            clean_folder_contents(item, recursive=False)

# =========================================================
# LÓGICA: IDENTIFICAÇÃO DE COLETORES
# =========================================================

def extract_collectors(cm_folder: Path):
    collectors = set()
    for file in cm_folder.iterdir():
        if not file.is_file(): continue
        df = read_table(file)
        if df.empty: continue

        unit_name_col = find_column(df, ["unit name"])
        unit_type_col = find_column(df, ["unit type"])

        if not unit_name_col or not unit_type_col: continue

        for _, row in df.iterrows():
            unit_type = str(row.get(unit_type_col, "")).lower()
            unit_name = str(row.get(unit_name_col, "")).strip()
            if "collector" in unit_type and unit_name:
                collectors.add(unit_name)
    return sorted(collectors)

# =========================================================
# LÓGICA: STAP STATUS (COM VERSÃO E FILTRO INATIVOS)
# =========================================================

def process_stap_status(folder: Path):
    records = []
    for file in folder.iterdir():
        if not file.is_file(): continue
        df = read_table(file)
        if df.empty: continue

        status_col = find_column(df, ["status"])
        host_col = find_column(df, ["software stap host", "stap host", "host"])
        ver_col = find_column(df, ["revision", "version", "s-tap revision", "stap revision"])

        if not status_col: continue

        for _, row in df.iterrows():
            status = str(row.get(status_col, "")).strip()
            host = str(row.get(host_col, "")).strip() if host_col else "N/A"
            version = str(row.get(ver_col, "")).strip() if ver_col else "Desc."
            
            if status:
                is_active = False
                status_lower = status.lower()
                if any(k in status_lower for k in ACTIVE_KEYWORDS):
                    is_active = True
                
                records.append({
                    "host": host,
                    "status": status,
                    "version": version,
                    "is_active": is_active
                })

    if not records:
        return pd.DataFrame(), {"active": 0, "inactive": 0, "total": 0}

    df_all = pd.DataFrame(records)
    
    active_count = df_all[df_all["is_active"] == True].shape[0]
    inactive_count = df_all[df_all["is_active"] == False].shape[0]

    return df_all, {
        "active": active_count,
        "inactive": inactive_count,
        "total": len(df_all)
    }

# =========================================================
# LÓGICA: AGREGAÇÃO (COM DATA)
# =========================================================

def analyze_aggregation_errors(base_folder: Path, collectors: list):
    """
    Retorna uma lista de dicionários com os erros encontrados:
    [{'collector': X, 'activity': Y, 'status': Z, 'date': D}, ...]
    """
    issues = []

    for collector in collectors:
        collector_path = base_folder / collector
        if not collector_path.exists():
            continue

        for file in collector_path.iterdir():
            if not file.is_file(): continue

            df = read_table(file)
            if df.empty: continue

            act_col = find_column(df, ["activity type", "activity", "process"])
            status_col = find_column(df, ["status", "execution status"])
            # Coluna de data (Start Time, Run Time, Timestamp)
            date_col = find_column(df, ["start time", "run time", "timestamp", "date"])

            if not act_col or not status_col:
                continue

            for _, row in df.iterrows():
                status_val = str(row.get(status_col, "")).strip()
                activity_val = str(row.get(act_col, "")).strip()
                
                date_val = str(row.get(date_col, "")).strip() if date_col else "Data desc."

                if not status_val or not activity_val:
                    continue

                is_success = any(ok_word in status_val.lower() for ok_word in SUCCESS_KEYWORDS)
                
                if not is_success:
                    issues.append({
                        "collector": collector,
                        "activity": activity_val,
                        "status": status_val,
                        "date": date_val
                    })

    try:
        issues.sort(key=lambda x: (x['collector'], x['date']), reverse=True)
    except:
        pass
        
    return issues

# =========================================================
# MAIN
# =========================================================

def main(base_path="."):
    # Ajusta o caminho base: se o script for executado de CM-Processor/, base_path deve ser '../' ou '.'
    # O argumento 'base_path' (que por padrão é '.') agora representa a raiz onde a pasta CM deve ser criada.
    base = Path(base_path).resolve()
    cm = base / "CM"
    
    print("--- INICIANDO GUARDIUM HEALTH CHECK ---")

    # 1. LIMPEZA
    print("🧹 Limpando arquivos de execuções anteriores...")
    clean_folder_contents(cm / "Central Management")
    clean_folder_contents(cm / "STAP status")
    clean_folder_contents(cm / "Processos de agregação", recursive=True)
    clean_folder_contents(cm / "Qualidade da coleta", recursive=True)
    
    # 2. CRIAÇÃO DE ESTRUTURA
    cm.mkdir(exist_ok=True)
    for sf in BASE_SUBFOLDERS:
        (cm / sf).mkdir(exist_ok=True)

    print("✔ Estrutura de pastas verificada.")
    input("\n➡ Coloque a planilha do Central Management na pasta 'Central Management' e pressione ENTER...")

    collectors = extract_collectors(cm / "Central Management")
    if not collectors:
        print("❌ Nenhum Collector encontrado.")
        return

    print(f"\n✔ {len(collectors)} Collectors identificados.")

    # 3. CRIAÇÃO DE SUBPASTAS
    for collector in collectors:
        (cm / "Processos de agregação" / collector).mkdir(parents=True, exist_ok=True)
        (cm / "Qualidade da coleta" / collector).mkdir(parents=True, exist_ok=True)

    input("\n➡ Agora coloque os arquivos nas subpastas (STAP status, Agregação/Collector X) e pressione ENTER...")

    # 4. PROCESSAMENTO E ANÁLISE
    print("\n⚙ Processando dados...")
    
    # A) STAP
    stap_df, stap_sum = process_stap_status(cm / "STAP status")
    
    # B) Agregação
    agg_issues = analyze_aggregation_errors(cm / "Processos de agregação", collectors)

    # 5. GERAÇÃO DE RELATÓRIO
    out = cm / "output"
    out.mkdir(exist_ok=True)

    # Excel (Dados Brutos)
    with pd.ExcelWriter(out / "CM_report.xlsx", engine="openpyxl") as writer:
        if not stap_df.empty:
            stap_df.to_excel(writer, sheet_name="STAP_Todos", index=False)
        if agg_issues:
            pd.DataFrame(agg_issues).to_excel(writer, sheet_name="Erros_Agregacao", index=False)

    # Word (Relatório Executivo)
    if Document:
        doc = Document()
        doc.add_heading("Relatório de Health Check - Guardium", 0)
        doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

        # --- SEÇÃO 1: STAP ---
        doc.add_heading("1. Status dos Agentes (STAP)", level=1)
        doc.add_paragraph(f"Total de agentes detectados: {stap_sum['total']}")
        doc.add_paragraph(f"Agentes ATIVOS: {stap_sum['active']}")
        p_inativos = doc.add_paragraph()
        p_inativos.add_run(f"Agentes INATIVOS: {stap_sum['inactive']}").bold = True

        # Tabela de Inativos (se houver)
        if stap_sum['inactive'] > 0:
            doc.add_heading("Detalhamento dos Agentes Inativos:", level=3)
            
            inativos_df = stap_df[stap_df["is_active"] == False]

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Host'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Versão (Revision)'

            for _, row in inativos_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['host'])
                row_cells[1].text = str(row['status'])
                row_cells[2].text = str(row['version'])
        else:
            doc.add_paragraph("✔ Todos os agentes reportados estão ativos.")

        # --- SEÇÃO 2: AGREGAÇÃO ---
        doc.add_heading("2. Falhas em Processos de Agregação", level=1)
        
        if agg_issues:
            doc.add_paragraph("Foram detectadas falhas nos seguintes processos (Purge, Export, Archive, etc.):")
            
            err_table = doc.add_table(rows=1, cols=3)
            err_table.style = 'Table Grid'
            eh_cells = err_table.rows[0].cells
            eh_cells[0].text = 'Coletor / Appliance'
            eh_cells[1].text = 'Falha (Processo/Status)'
            eh_cells[2].text = 'Data Ocorrência'

            for issue in agg_issues:
                row_cells = err_table.add_row().cells
                row_cells[0].text = str(issue['collector'])
                row_cells[1].text = f"{issue['activity']} ({issue['status']})"
                row_cells[2].text = str(issue['date'])
        else:
            doc.add_paragraph("Nenhum erro crítico encontrado nos logs de agregação fornecidos.")

        word_path = out / "Relatorio_Executivo.docx"
        doc.save(word_path)
        print(f"📄 Relatório Word gerado: {word_path}")

    print("\n✅ PROCESSAMENTO FINALIZADO")
    print(f"Agentes Inativos: {stap_sum['inactive']}")
    print(f"Erros de Agregação: {len(agg_issues)}")

if __name__ == "__main__":
    arg = sys.argv[1] if len(sys.argv) > 1 else "." 
    main(arg)
